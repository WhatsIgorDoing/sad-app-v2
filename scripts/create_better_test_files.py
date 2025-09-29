"""
Script para criar arquivos de teste melhores para o ProfiledExtractorService
"""

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_pdf_test_file():
    """Cria um arquivo PDF de teste mais robusto"""
    pdf_path = Path("tests/fixtures/documento_rir.pdf")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # Criar PDF com ReportLab
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter

    # Título
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 100, "Relatório de Inspeção Robótica")

    # Código do documento
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, height - 150, "Relatório: CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A")

    # Conteúdo
    c.setFont("Helvetica", 12)
    y_position = height - 200

    content = [
        "1. INTRODUÇÃO",
        "",
        "Este relatório apresenta os resultados da inspeção robótica realizada",
        "no equipamento CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A.",
        "",
        "2. METODOLOGIA",
        "",
        "A inspeção foi realizada utilizando robô especializado seguindo",
        "procedimentos padrão da indústria.",
        "",
        "3. RESULTADOS",
        "",
        "Os resultados da inspeção indicam condições satisfatórias do",
        "equipamento avaliado.",
    ]

    for line in content:
        c.drawString(100, y_position, line)
        y_position -= 20

    c.save()
    print(f"PDF criado: {pdf_path}")
    return pdf_path


def create_docx_test_file():
    """Cria um arquivo DOCX de teste"""
    docx_path = Path("tests/fixtures/documento_rir.docx")
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Título
    doc.add_heading("Relatório de Inspeção Robótica", 0)

    # Código do documento
    doc.add_paragraph(
        "Relatório: CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A", style="Heading 2"
    )

    # Conteúdo
    doc.add_heading("1. INTRODUÇÃO", level=1)
    doc.add_paragraph(
        "Este relatório apresenta os resultados da inspeção robótica realizada "
        "no equipamento CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A."
    )

    doc.add_heading("2. METODOLOGIA", level=1)
    doc.add_paragraph(
        "A inspeção foi realizada utilizando robô especializado seguindo "
        "procedimentos padrão da indústria."
    )

    doc.add_heading("3. RESULTADOS", level=1)
    doc.add_paragraph(
        "Os resultados da inspeção indicam condições satisfatórias do "
        "equipamento avaliado."
    )

    doc.save(str(docx_path))
    print(f"DOCX criado: {docx_path}")
    return docx_path


def create_pid_test_files():
    """Cria arquivos de teste para perfil PID"""
    # PDF PID
    pdf_path = Path("tests/fixtures/documento_pid.pdf")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 100, "P&ID - Diagrama de Processo")

    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, height - 150, "Documento: PID-CZ6-RNEST-U22-3.2.1_REV-A")

    c.setFont("Helvetica", 12)
    c.drawString(
        100, height - 200, "Este P&ID representa o sistema de processo da unidade."
    )

    c.save()

    # DOCX PID
    docx_path = Path("tests/fixtures/documento_pid.docx")
    doc = Document()
    doc.add_heading("P&ID - Diagrama de Processo", 0)
    doc.add_paragraph("Documento: PID-CZ6-RNEST-U22-3.2.1_REV-A", style="Heading 2")
    doc.add_paragraph("Este P&ID representa o sistema de processo da unidade.")
    doc.save(str(docx_path))

    print(f"Arquivos PID criados: {pdf_path} e {docx_path}")


if __name__ == "__main__":
    try:
        create_pdf_test_file()
        create_docx_test_file()
        create_pid_test_files()
        print("Todos os arquivos de teste foram criados com sucesso!")
    except Exception as e:
        print(f"Erro ao criar arquivos: {e}")
